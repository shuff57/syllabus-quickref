#!/usr/bin/env python3
"""Extract plain text from a syllabus file so the rest of the pipeline has a
single source of truth to check the rewrite against.

Usage:
    python extract.py INPUT [-o OUT.txt]

Handles .pdf (PyMuPDF, falls back to pypdf), .docx (python-docx, tables
included), and .txt/.md (passthrough). Writes UTF-8 and prints the output path.

Why this exists: the checker compares every number in the quick reference back
against the source text. That comparison is only meaningful if the source text
is a faithful, complete dump, including tables, where grading weights live.
"""
import argparse
import os
import sys


def from_pdf(path):
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
    except ImportError:
        from pypdf import PdfReader
        return "\n".join((p.extract_text() or "") for p in PdfReader(path).pages)


def from_docx(path):
    from docx import Document
    doc = Document(path)
    out = [p.text for p in doc.paragraphs]
    for t in doc.tables:                       # grading weights usually live here
        for row in t.rows:
            out.append("\t".join(c.text.strip() for c in row.cells))
    return "\n".join(out)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("-o", "--out")
    a = ap.parse_args()

    ext = os.path.splitext(a.input)[1].lower()
    if ext == ".pdf":
        text = from_pdf(a.input)
    elif ext in (".docx", ".doc"):
        text = from_docx(a.input)
    elif ext in (".txt", ".md", ".markdown", ""):
        with open(a.input, encoding="utf-8", errors="replace") as f:
            text = f.read()
    else:
        sys.exit("[FAIL] unsupported extension %r. Convert to pdf/docx/txt first" % ext)

    out = a.out or os.path.splitext(a.input)[0] + ".source.txt"
    with open(out, "w", encoding="utf-8") as f:
        f.write(text)

    words = len(text.split())
    print("[OK] %d words -> %s" % (words, out))
    if words < 100:
        print("[WARN] very little text extracted. If the PDF is a scan, read it "
              "with the Read tool instead and save the text by hand.")


if __name__ == "__main__":
    main()
