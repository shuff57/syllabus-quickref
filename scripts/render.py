#!/usr/bin/env python3
"""Render a syllabus quick reference into the formats a class actually needs.

Usage:
    python render.py QUICKREF.md [--html OUT.html] [--docx OUT.docx] [--columns 2]
                                 [--theme NAME] [--sticky "Head|body"]
    python render.py --list-themes

Themes live in scripts/themes/NAME.css and only have to redefine the token set
at the top of spec.css; diagrams.py paints the weight bar from those same names,
so a theme recolours the figure for free. A theme needing an SVG filter ships a
NAME.defs.html beside it, injected at the top of <body>.

Reads a deliberately small slice of Markdown: h1, h2, bullets, pipe tables,
bold spans, and > for the closing note. A quick reference should not need
anything more expressive than that, and a rewrite that wants a feature this does
not support is too complicated for a one-pager.

No third-party Markdown dependency on purpose: this runs on a school laptop
with whatever Python happens to be installed.
"""
import argparse
import html
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import diagrams

# bookSHelf house theme: parchment canvas, one Wedgwood accent held under 5% of
# the surface, warm ink ramp, serif headings capped at weight 500, no italic.
# Full spec: bookSHelf/.claude/skills/theme-factory/themes/bookshelf.md
CSS = """
@page { size: letter; margin: 0; }
:root {
  --wedgwood: #4e6e8e; --wedgwood-deep: #3d5a80;
  --parchment: #f5f4ed; --ivory: #faf9f5; --warm-sand: #e8e6dc;
  --ink: #141413; --charcoal: #3d3d3a; --olive: #504e49; --stone: #6b6a64;
  --border-cream: #f0eee6; --border-warm: #e8e6dc; --ring-warm: #d1cfc5;
  --serif: "Source Serif 4", "Source Serif Pro", Charter, Georgia, "Times New Roman", serif;
  --sans: Inter, system-ui, -apple-system, "Segoe UI", Helvetica, Arial, sans-serif;
}
* { box-sizing: border-box; }
body { font: 10.5pt/1.55 var(--sans); color: var(--ink);
       background: var(--parchment); margin: 0 auto; padding: 0.45in; max-width: 8.5in;
       -webkit-print-color-adjust: exact; print-color-adjust: exact; }
h1 { font: 500 25pt/1.1 var(--serif); margin: 0 0 3pt; letter-spacing: -0.4px; }
.sub { font-size: 10pt; color: var(--olive); margin: 0 0 10pt;
       padding-bottom: 7pt; border-bottom: 1px solid var(--wedgwood); }
.cols { column-count: __COLS__; column-gap: 26pt; counter-reset: sec; }
/* Sections never split across the column boundary: a rule you are scanning for
   should be whole and under its own heading, not continued overleaf. That makes
   section length the thing that has to be managed. Keep each one to roughly
   four or five bullets and the columns pack evenly on their own. One nine-item
   section cannot pack, and it is the section that wants dividing anyway. */
section { break-inside: avoid; margin: 0 0 10pt; counter-increment: sec; }
h2 { font: 500 13pt/1.2 var(--serif); color: var(--ink); margin: 0 0 5pt;
     padding-bottom: 3pt; border-bottom: 1px solid var(--border-warm);
     break-after: avoid; break-inside: avoid; }
h2::before { content: counter(sec, decimal-leading-zero);
             font: 500 9pt/1.4 var(--serif); color: var(--wedgwood);
             letter-spacing: 0.5px; margin-right: 6pt; }
p { margin: 0 0 5pt; color: var(--charcoal); }
ul { margin: 0; padding: 0; list-style: none; }
li { position: relative; padding-left: 13pt; margin-bottom: 2.5pt;
     color: var(--charcoal); break-inside: avoid; }
table, tr { break-inside: avoid; }
svg.fig { display: block; width: 100%; height: auto; margin: 3pt 0 5pt;
          break-inside: avoid; }
li::before { content: "\\2013"; position: absolute; left: 0;
             color: var(--wedgwood); }
table { border-collapse: collapse; width: 100%; margin: 2pt 0 4pt;
        font-variant-numeric: tabular-nums; }
th { text-align: left; font: 500 8pt/1.4 var(--sans); text-transform: uppercase;
     letter-spacing: 1.2px; color: var(--charcoal); background: var(--warm-sand);
     padding: 2pt 6pt; }
td { padding: 2pt 6pt; border-bottom: 1px solid var(--border-cream);
     color: var(--charcoal); }
strong { font-weight: 600; color: var(--ink); }
a { color: var(--wedgwood-deep); text-decoration: none; }
.note { column-span: all; margin-top: 6pt; padding-top: 5pt;
        border-top: 1px solid var(--border-warm); font-size: 9pt;
        color: var(--stone); }
/* The who/where/when strip is emitted as one span per '·' segment so grid-based
   themes can render it as spec cells. This theme keeps it as running text, so
   the separators have to come back. */
.sub span + span::before { content: "\\00b7"; margin: 0 6pt; color: var(--ring-warm); }
.sub em { font-style: normal; color: var(--stone); margin-right: 4pt; }
.head { display: block; }
.sticky { display: none; }   /* no hand in this theme; the note has nowhere to go */
/* The three hand marks, held to this theme's one accent rather than a marker set. */
.hi { background: #f3e7c0; padding: 0 .18em; border-radius: 2px; color: var(--ink);
      font-weight: 600; }
.circ { border: 1px solid var(--wedgwood); border-radius: 10px; padding: 0 5pt;
        color: var(--wedgwood-deep); font-weight: 600; }
.hand { font-style: italic; color: var(--wedgwood-deep); }
@media print { .note { page-break-inside: avoid; } }
"""

THEME_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "themes")

# Appended to every theme, so the button looks native in all of them and there
# is only one copy to fix. Colours fall back, because a theme is only obliged to
# define the token set diagrams.py needs.
EXPORT_CSS = """
/* The button rides on the closing note's own line, so it needs no rule or
   spacing of its own and cannot add height to the sheet. Each theme already
   styles .note; only the layout is added here. */
.note { display: flex; align-items: center; justify-content: space-between;
        gap: 12pt; }
.export-btn {
  font: 600 8pt/1 var(--sans, system-ui, -apple-system, "Segoe UI", sans-serif);
  letter-spacing: .08em; text-transform: uppercase; white-space: nowrap;
  color: var(--ink, #1a1a1a); background: var(--white, #fff);
  border: 1px solid var(--ink, #1a1a1a); border-radius: 3px;
  padding: 6pt 11pt; cursor: pointer; flex: 0 0 auto;
}
.export-btn:hover { background: var(--ink, #1a1a1a); color: var(--white, #fff); }
.export-btn:focus-visible { outline: 2px solid var(--wedgwood, #4e6e8e);
                            outline-offset: 2px; }
/* The button is not part of the handout. */
@media print { .export-btn { display: none !important; } }
"""

EXPORT_BTN = ('<button type="button" id="export-pdf" class="export-btn">'
              'Save as PDF</button>')

EXPORT_JS = """
<script>
document.getElementById("export-pdf").addEventListener("click", function () {
  window.print();
});
</script>
"""


def themes():
    """-> sorted names of the installed themes, 'bookshelf' always first."""
    found = []
    if os.path.isdir(THEME_DIR):
        found = sorted(f[:-4] for f in os.listdir(THEME_DIR)
                       if f.endswith(".css") and not f.startswith("_"))
    return ["bookshelf"] + [t for t in found if t != "bookshelf"]


def load_theme(name):
    """-> (css, defs_html). 'bookshelf' is the built-in above."""
    if name == "bookshelf":
        return CSS, ""
    css_path = os.path.join(THEME_DIR, name + ".css")
    if not os.path.exists(css_path):
        sys.exit("[FAIL] no theme %r. Installed: %s" % (name, ", ".join(themes())))
    defs_path = os.path.join(THEME_DIR, name + ".defs.html")
    defs = ""
    if os.path.exists(defs_path):
        defs = open(defs_path, encoding="utf-8").read()
    return open(css_path, encoding="utf-8").read(), defs


#  ==swiped==    a highlighter band behind the words
#  ((circled))   a pen circle round a hard rule
#  ~~in hand~~   a handwritten aside
# Three marks, because those are the three a teacher actually makes on a page.
# Themes that have no hand simply render them as emphasis; nothing is lost.
#
# All three delimiters are two characters. A single `~` was tried first and is
# wrong: one stray tilde anywhere in a syllabus silently turns the rest of the
# line into handwriting, and nothing downstream would catch it. Doubling costs
# one character and removes the failure. GFM reads `~~` as strikethrough, which
# this dialect does not have, so nothing is being shadowed.
MARKS = [
    (re.compile(r"==(.+?)=="), r'<span class="hi">\1</span>'),
    (re.compile(r"\(\((.+?)\)\)"), r'<span class="circ">\1</span>'),
    (re.compile(r"~~(.+?)~~"), r'<span class="hand">\1</span>'),
]
STRIP_MARKS = re.compile(r"==|\(\(|\)\)|~~")

# One pass over a line, in the order the marks appear, so the .docx can rebuild
# it as Word runs. Groups: 1 bold, 2 swiped, 3 circled, 4 in hand.
TOKEN = re.compile(r"\*\*(.+?)\*\*|==(.+?)==|\(\((.+?)\)\)|~~(.+?)~~")


def lint_marks(md):
    """-> [warning]. An unclosed mark renders as literal '==' on a handout and
    nothing else complains, so say so here rather than at the printer."""
    out = []
    for n, line in enumerate(md.splitlines(), 1):
        for delim, count in (("==", line.count("==")), ("~~", line.count("~~"))):
            if count % 2:
                out.append("line %d: odd number of %s" % (n, delim))
        if line.count("((") != line.count("))"):
            out.append("line %d: (( and )) do not pair" % n)
    return out


def inline(s):
    s = html.escape(s)
    s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
    for pat, rep in MARKS:
        s = pat.sub(rep, s)
    s = re.sub(r"(?<![\">])(https?://[^\s<)]+)", r'<a href="\1">\1</a>', s)
    return s


def parse(md):
    """-> (title, subtitle, [(heading, [blocks])], [notes]).

    A block is ('ul', [items]) | ('table', [rows]) | ('p', text).
    """
    title, sub, notes, sections = "", "", [], []
    cur = None
    lines = md.splitlines()
    i = 0
    # A blank line ends the current list or table. Without this, two tables in
    # one section fuse into one: a grading table followed by a grade-scale table
    # became a single eight-column block that no longer looked like weights, so
    # the figure silently disappeared and a merged table took its place.
    fresh = True
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1
        if not line:
            fresh = True
            continue
        if line.startswith("```flow"):          # a sequence to draw as a chain
            steps = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                if lines[i].strip():
                    steps.append(lines[i].strip().lstrip("- "))
                i += 1
            i += 1
            if cur is None:
                cur = ("", [])
                sections.append(cur)
            cur[1].append(("flow", steps))
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            for j in range(i, len(lines)):          # first non-blank line after the
                nxt = lines[j].strip()              # title is the who/where/when strip
                if not nxt:
                    continue
                if not nxt.startswith(("#", "-", "*", "|", ">")):
                    sub, i = nxt, j + 1
                break
            continue
        if line.startswith("## "):
            cur = (line[3:].strip(), [])
            sections.append(cur)
            continue
        if line.startswith(">"):
            notes.append(line.lstrip("> ").strip())
            continue
        if cur is None:
            cur = ("", [])
            sections.append(cur)
        if line.startswith(("- ", "* ")):
            if cur[1] and cur[1][-1][0] == "ul" and not fresh:
                cur[1][-1][1].append(line[2:].strip())
            else:
                cur[1].append(("ul", [line[2:].strip()]))
        elif line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                continue                                  # separator row
            if cur[1] and cur[1][-1][0] == "table" and not fresh:
                cur[1][-1][1].append(cells)
            else:
                cur[1].append(("table", [cells]))
        else:
            cur[1].append(("p", line))
        fresh = False
    return title, sub, sections, notes


def to_html(md, columns, css=None, defs="", sticky=None):
    title, sub, sections, notes = parse(md)
    out = [(css or CSS).replace("__COLS__", str(columns))]
    body = [defs] if defs else []
    body.append('<header class="head"><div class="headline">')
    body.append('<h1>%s</h1>' % inline(title))
    if sub:
        # One span per '·' segment: running text in some themes, spec cells in
        # others. A segment written "Label: value" carries its label in an <em>,
        # which grid themes stack above the value and the rest print inline.
        cells = []
        for p in re.split(r"\s*[··]\s*", sub):
            p = p.strip()
            if not p:
                continue
            label, sep, val = p.partition(": ")
            cells.append("<span><em>%s</em>%s</span>" % (inline(label), inline(val))
                         if sep else "<span>%s</span>" % inline(p))
        body.append('<p class="sub">%s</p>' % "".join(cells))
    body.append('</div>')
    if sticky:
        head, _, rest = sticky.partition("|")
        body.append('<div class="sticky"><b>%s</b>%s</div>'
                    % (inline(head.strip()), inline(rest.strip())))
    body.append('</header>')
    body.append('<div class="cols">')
    for head, blocks in sections:
        body.append("<section>")
        if head:
            body.append("<h2>%s</h2>" % inline(head))
        for kind, payload in blocks:
            if kind == "ul":
                body.append("<ul>%s</ul>" %
                            "".join("<li>%s</li>" % inline(x) for x in payload))
            elif kind == "flow":
                body.append(diagrams.flow(payload) or
                            "<ul>%s</ul>" % "".join("<li>%s</li>" % inline(x)
                                                    for x in payload))
            elif kind == "table":
                fig = diagrams.weight_bar(payload)
                if fig:
                    body.append(fig)
                    continue
                head_row, *rest = payload
                body.append("<table><tr>%s</tr>%s</table>" % (
                    "".join("<th>%s</th>" % inline(c) for c in head_row),
                    "".join("<tr>%s</tr>" % "".join("<td>%s</td>" % inline(c) for c in r)
                            for r in rest)))
            else:
                body.append("<p>%s</p>" % inline(payload))
        body.append("</section>")
    body.append("</div>")
    # the button sits on the last note's line; with no note it gets its own
    if notes:
        for i, n in enumerate(notes):
            body.append('<p class="note"><span>%s</span>%s</p>'
                        % (inline(n), EXPORT_BTN if i == len(notes) - 1 else ""))
    else:
        body.append('<p class="note"><span></span>%s</p>' % EXPORT_BTN)
    body.append(EXPORT_JS)
    return ("<!doctype html><html><head><meta charset=\"utf-8\">"
            "<title>%s</title><style>%s</style></head><body>%s</body></html>"
            % (html.escape(title), out[0] + EXPORT_CSS, "".join(body)))


def theme_palette(name):
    """Pull the colours the .docx needs out of the theme's own CSS, so the Word
    file and the page cannot drift apart. Unset tokens keep the bookshelf ramp."""
    pal = {"ink": "141413", "charcoal": "3D3D3A", "stone": "6B6A64",
           "accent": "4E6E8E", "accent_deep": "3D5A80", "line": "D1CFC5",
           "block": "E8E6DC", "white": "FAF9F5", "hi": "FFE08A"}
    css_path = os.path.join(THEME_DIR, name + ".css")
    if not os.path.exists(css_path):
        return pal
    css = open(css_path, encoding="utf-8").read()
    for token, key in (("ink", "ink"), ("charcoal", "charcoal"), ("stone", "stone"),
                       ("wedgwood", "accent"), ("wedgwood-deep", "accent_deep"),
                       ("line", "line"), ("block", "block"), ("white", "white"),
                       ("hi", "hi")):
        m = re.search(r"--%s\s*:\s*#([0-9A-Fa-f]{6})\b" % re.escape(token), css)
        if m:
            pal[key] = m.group(1).upper()
    return pal


BROWSERS = [
    r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
    r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
    "/Applications/Microsoft Edge.app/Contents/MacOS/Microsoft Edge",
    "google-chrome", "chromium", "chromium-browser", "microsoft-edge",
]


def to_pdf(html_path, pdf_path):
    """Print the HTML through a browser, which is the only renderer that agrees
    with the page exactly — same engine, same CSS, same @page rules.

    The .docx rebuilds the layout out of Word primitives and is a good likeness;
    this is the original. Use it for the copies that get handed out, and the
    .docx for the copy someone needs to edit."""
    import shutil
    import subprocess

    exe = next((b for b in BROWSERS
                if os.path.exists(b) or shutil.which(b)), None)
    if not exe:
        return ("no Chrome or Edge found. Open the .html and print it "
                "(Ctrl+P, Save as PDF) — the result is identical.")
    src = "file:///" + os.path.abspath(html_path).replace("\\", "/")
    cmd = [shutil.which(exe) or exe, "--headless=new", "--disable-gpu",
           "--no-pdf-header-footer",
           "--print-to-pdf=" + os.path.abspath(pdf_path), src]
    try:
        subprocess.run(cmd, capture_output=True, timeout=120)
    except Exception as e:                      # noqa: BLE001 - report, don't crash
        return "browser failed: %s" % e
    if not os.path.exists(pdf_path):
        return "browser produced no file. Print the .html by hand instead."
    return None


PCT = re.compile(r"^\s*(\d+(?:\.\d+)?)\s*%\s*$")


def weight_rows(payload):
    """-> [(label, percent)] when a table is a set of weights summing to 100,
    else None. Same test diagrams.py uses to decide the page gets a bar."""
    head, *rest = payload
    if len(head) != 2 or len(rest) < 2:
        return None
    out = []
    for row in rest:
        if len(row) != 2:
            return None
        m = PCT.match(row[1])
        if not m:
            return None
        out.append((re.sub(r"\*\*", "", row[0]).strip(), float(m.group(1))))
    return out if abs(sum(v for _, v in out) - 100) < 0.51 else None


def to_docx(md, path, theme="bookshelf", columns=2, sticky=None):
    """Rebuild the page in Word: two columns, one bordered compartment per
    section, the spec bar, the corner note.

    Word is a flow document, so this is a reconstruction rather than a
    conversion — a bordered box is a one-column table, a column is a section
    property. The point is that a teacher who opens the .docx recognises the
    handout they printed, and can still edit every word of it."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_BREAK
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pal = theme_palette(theme)
    rgb = lambda h: RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    INK, CHARCOAL, STONE = rgb(pal["ink"]), rgb(pal["charcoal"]), rgb(pal["stone"])
    ACCENT, ACCENT_DEEP = rgb(pal["accent"]), rgb(pal["accent_deep"])
    SERIF = "Georgia"       # local stand-in for Bodoni / Source Serif
    SANS = "Segoe UI"       # local stand-in for Inter / DM Sans
    HAND = "Segoe Script"   # local stand-in for Caveat, as in the CSS

    # ---- low-level Word plumbing -------------------------------------------
    def shade(el, fill):
        e = OxmlElement("w:shd")
        e.set(qn("w:val"), "clear")
        e.set(qn("w:fill"), fill)
        el.append(e)

    def cell_shade(cell, fill):
        shade(cell._tc.get_or_add_tcPr(), fill)

    def table_borders(tbl, color, sz=4, val="single", inside=True):
        b = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement("w:" + edge)
            on = val if (inside or not edge.startswith("inside")) else "nil"
            e.set(qn("w:val"), on)
            e.set(qn("w:sz"), str(sz if on != "nil" else 0))
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), color if on != "nil" else "auto")
            b.append(e)
        tbl._tbl.tblPr.append(b)

    def rule(par, color, sz=18):
        """A thick rule under a paragraph — the masthead's hard line."""
        pPr = par._p.get_or_add_pPr()
        b = OxmlElement("w:pBdr")
        e = OxmlElement("w:bottom")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "2")
        e.set(qn("w:color"), color)
        b.append(e)
        pPr.append(b)

    def set_columns(section, n):
        cols = section._sectPr.xpath("./w:cols")[0]
        cols.set(qn("w:num"), str(n))
        cols.set(qn("w:space"), "260")          # ~0.18in gutter

    def tight(par, before=0, after=0, spacing=1.15):
        f = par.paragraph_format
        f.space_before, f.space_after = Pt(before), Pt(after)
        f.line_spacing = spacing
        return par

    def boxed(run):
        """Word cannot draw an ellipse behind text, so ((circled)) becomes a
        character border — the nearest thing Word has that still reads as
        'someone ringed this'."""
        bdr = OxmlElement("w:bdr")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "6")
        bdr.set(qn("w:space"), "0")
        bdr.set(qn("w:color"), pal["accent"])
        run._element.get_or_add_rPr().append(bdr)

    def add_runs(p, text, size=None, color=None):
        """Split text into runs carrying **bold** and the three hand marks.
        Word has a real highlighter and real fonts, so ==swiped== and ~~in
        hand~~ survive as themselves."""
        pos = 0
        for m in TOKEN.finditer(text):
            if m.start() > pos:
                p.add_run(text[pos:m.start()])
            if m.group(1) is not None:
                r = p.add_run(m.group(1)); r.bold = True; r.font.color.rgb = INK
            elif m.group(2) is not None:
                r = p.add_run(m.group(2)); r.bold = True
                r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                r.font.color.rgb = INK
            elif m.group(3) is not None:
                r = p.add_run(m.group(3)); r.bold = True
                r.font.color.rgb = ACCENT; boxed(r)
            else:
                r = p.add_run(m.group(4))
                r.font.name = HAND; r.font.color.rgb = ACCENT
            pos = m.end()
        if pos < len(text):
            p.add_run(text[pos:])
        for r in p.runs:
            if size is not None:
                r.font.size = size
            if color is not None and r.font.color.rgb is None:
                r.font.color.rgb = color
        return p

    # ---- document ----------------------------------------------------------
    title, sub, sections, notes = parse(md)
    doc = Document()
    s0 = doc.sections[0]
    s0.top_margin = s0.bottom_margin = Inches(0.4)
    s0.left_margin = s0.right_margin = Inches(0.4)
    PAGE_W = 7.7                                   # 8.5in less both margins

    normal = doc.styles["Normal"]
    normal.font.name = SANS
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = CHARCOAL
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(2)

    # masthead: title, and the corner note beside it rather than over it
    mast = doc.add_table(rows=1, cols=2 if sticky else 1)
    mast.autofit = False
    table_borders(mast, "auto", 0, "nil")
    left = mast.rows[0].cells[0]
    left.width = Inches(PAGE_W - 1.6 if sticky else PAGE_W)
    left.paragraphs[0].text = ""
    r = tight(left.paragraphs[0], spacing=0.95).add_run(re.sub(r"\*\*", "", title))
    r.font.name, r.font.size, r.font.color.rgb = SERIF, Pt(26), INK
    if sticky:
        head, _, rest = sticky.partition("|")
        note = mast.rows[0].cells[1]
        note.width = Inches(1.6)
        cell_shade(note, pal["hi"])
        note.paragraphs[0].text = ""
        r = tight(note.paragraphs[0], before=3, spacing=0.95).add_run(head.strip())
        r.font.name, r.font.size, r.font.bold = HAND, Pt(12), True
        r.font.color.rgb = ACCENT
        p = tight(note.add_paragraph(), after=3, spacing=0.95)
        r = p.add_run(rest.strip())
        r.font.name, r.font.size, r.font.color.rgb = HAND, Pt(10), INK

    # spec bar: label stacked over value, one cell per '·' segment
    if sub:
        cells = [c.strip() for c in re.split(r"\s*[··]\s*", sub) if c.strip()]
        bar = doc.add_table(rows=1, cols=len(cells))
        bar.autofit = False
        table_borders(bar, pal["line"], 4)
        for cell, text in zip(bar.rows[0].cells, cells):
            cell.width = Inches(PAGE_W / len(cells))
            cell_shade(cell, pal["white"])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            label, sep, val = text.partition(": ")
            cell.paragraphs[0].text = ""
            if sep:
                r = tight(cell.paragraphs[0]).add_run(label.upper())
                r.font.name, r.font.size = SANS, Pt(6.5)
                r.font.color.rgb = STONE
                p = tight(cell.add_paragraph())
            else:
                val, p = text, tight(cell.paragraphs[0])
            r = p.add_run(val)
            r.font.name, r.font.size, r.font.bold = SERIF, Pt(10.5), True
            r.font.color.rgb = INK

    rule(tight(doc.add_paragraph(), before=4, after=6), pal["ink"], 18)

    # everything below the masthead runs in columns
    body = doc.add_section(WD_SECTION.CONTINUOUS)
    body.top_margin = body.bottom_margin = Inches(0.4)
    body.left_margin = body.right_margin = Inches(0.4)
    set_columns(body, columns)
    COL_W = (PAGE_W - 0.18 * (columns - 1)) / columns

    for n, (head, blocks) in enumerate(sections, 1):
        box = doc.add_table(rows=1, cols=1)
        box.autofit = False
        table_borders(box, pal["line"], 4)
        cell = box.rows[0].cells[0]
        cell.width = Inches(COL_W)

        first = cell.paragraphs[0]
        first.text = ""
        if head:
            # the compartment header bar: shaded, with the coded id
            hp = tight(first, before=1, after=3)
            shade(hp._p.get_or_add_pPr(), pal["block"])
            r = hp.add_run("%02d   " % n)
            r.font.name, r.font.size, r.font.bold = SANS, Pt(7), True
            r.font.color.rgb = ACCENT
            r = hp.add_run(head)
            r.font.name, r.font.size, r.font.color.rgb = SERIF, Pt(12), INK
            first = None

        for kind, payload in blocks:
            if kind == "flow":
                p = tight(cell.add_paragraph() if first is None else first)
                first = None
                add_runs(p, "  →  ".join(payload))
            elif kind == "ul":
                for item in payload:
                    p = cell.add_paragraph() if first is None else first
                    first = None
                    tight(p, after=1.5)
                    r = p.add_run("–  ")
                    r.font.color.rgb = ACCENT
                    add_runs(p, item)
                    p.paragraph_format.left_indent = Inches(0.14)
                    p.paragraph_format.first_line_indent = Inches(-0.14)
            elif kind == "table":
                weights = weight_rows(payload)
                if weights:
                    # The proportional bar, built out of table cells rather than
                    # the page's SVG. Word cannot place an SVG, and rasterising
                    # one would put a picture of a fact into a file whose whole
                    # job is being editable. Cell widths carry the proportion,
                    # so the 75% block still reads as three quarters at a glance.
                    order = sorted(range(len(weights)), key=lambda i: -weights[i][1])
                    fills = {order[0]: pal["accent"], order[1]: pal["block"]}
                    bar = cell.add_table(rows=1, cols=len(weights))
                    bar.autofit = False
                    table_borders(bar, pal["line"], 4)
                    for i, (lab, v) in enumerate(weights):
                        c = bar.rows[0].cells[i]
                        c.width = Inches(max(COL_W * v / 100.0, 0.3))
                        cell_shade(c, fills.get(i, pal["white"]))
                        c.text = ""
                        p = tight(c.paragraphs[0])
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run("%g%%" % v)
                        r.font.name, r.font.size, r.font.bold = SANS, Pt(8), True
                        r.font.color.rgb = rgb(pal["white"]) if i == order[0] else INK
                    for i, (lab, v) in enumerate(weights):
                        p = tight(cell.add_paragraph(), after=0.5)
                        r = p.add_run("■  ")
                        r.font.size = Pt(7)
                        r.font.color.rgb = rgb(fills.get(i, pal["line"]))
                        r = p.add_run(lab)
                        r.font.size, r.font.color.rgb = Pt(8.5), CHARCOAL
                        r = p.add_run("   %g%%" % v)
                        r.font.size, r.font.bold, r.font.color.rgb = Pt(8.5), True, INK
                    continue
                head_row, *rest = payload
                t = cell.add_table(rows=1, cols=len(head_row))
                table_borders(t, pal["line"], 4)
                for c, txt in zip(t.rows[0].cells, head_row):
                    cell_shade(c, pal["block"])
                    c.text = ""
                    r = tight(c.paragraphs[0]).add_run(txt.upper())
                    r.font.name, r.font.size = SANS, Pt(6.5)
                    r.font.color.rgb = STONE
                for row in rest:
                    for c, txt in zip(t.add_row().cells, row):
                        c.text = ""
                        add_runs(tight(c.paragraphs[0]), txt,
                                 size=Pt(9), color=CHARCOAL)
            else:
                p = tight(cell.add_paragraph() if first is None else first)
                first = None
                add_runs(p, payload)

        tight(doc.add_paragraph(), after=0, spacing=0.5)   # gap between boxes
        # Word fills column one to the bottom before starting column two, where
        # CSS balances them. An explicit break at the halfway section restores
        # the page's split instead of leaving column two nearly empty.
        if columns > 1 and n == (len(sections) + 1) // 2:
            doc.add_paragraph().add_run().add_break(WD_BREAK.COLUMN)

    for note_text in notes:
        p = tight(doc.add_paragraph(), before=6)
        add_runs(p, note_text, size=Pt(8), color=STONE)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(path)



def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("quickref", nargs="?")
    ap.add_argument("--html")
    ap.add_argument("--docx")
    ap.add_argument("--pdf", help="print the HTML through headless Chrome/Edge")
    ap.add_argument("--columns", type=int, default=2)
    ap.add_argument("--theme", default="bookshelf",
                    help="one of: " + ", ".join(themes()))
    ap.add_argument("--sticky", metavar='"Head|body"',
                    help="corner note; themes that support it place it top right")
    ap.add_argument("--list-themes", action="store_true")
    a = ap.parse_args()
    if a.list_themes:
        for t in themes():
            print(t)
        return
    if not a.quickref:
        ap.error("quickref is required (or pass --list-themes)")
    md = open(a.quickref, encoding="utf-8").read()
    for w in lint_marks(md):
        print("[WARN] unclosed mark, " + w, file=sys.stderr)
    if not a.html and not a.docx and not a.pdf:
        a.html = a.quickref.rsplit(".", 1)[0] + ".html"
    # a PDF is printed FROM the html, so one has to exist to print
    if a.pdf and not a.html:
        a.html = a.pdf.rsplit(".", 1)[0] + ".html"
    if a.html:
        css, defs = load_theme(a.theme)
        open(a.html, "w", encoding="utf-8").write(
            to_html(md, a.columns, css, defs, a.sticky))
        print("[OK] wrote %s  (theme: %s)" % (a.html, a.theme))
    if a.pdf:
        err = to_pdf(a.html, a.pdf)
        print("[FAIL] " + err if err else "[OK] wrote " + a.pdf)
    if a.docx:
        to_docx(md, a.docx, a.theme, a.columns, a.sticky)
        print("[OK] wrote %s  (theme: %s)" % (a.docx, a.theme))


if __name__ == "__main__":
    main()
